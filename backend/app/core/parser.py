"""文档解析模块：支持 PDF / Word / TXT / Markdown"""
from __future__ import annotations

from pathlib import Path
from typing import Literal

from .llm import LLMError

SupportedExt = Literal[".pdf", ".docx", ".doc", ".txt", ".md"]


class ParseError(Exception):
    pass


def parse_file(file_path: Path, max_chars: int = 50000) -> str:
    """根据扩展名分派到具体解析器"""
    ext = file_path.suffix.lower()
    try:
        if ext == ".pdf":
            text = _parse_pdf(file_path)
        elif ext == ".docx":
            text = _parse_docx(file_path)
        elif ext == ".txt" or ext == ".md":
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        elif ext == ".doc":
            # 老 doc 格式不支持，提示转换
            raise ParseError(".doc 格式不支持，请转换为 .docx 或 .pdf 后上传")
        else:
            raise ParseError(f"不支持的文件格式: {ext}")
    except ParseError:
        raise
    except Exception as e:
        raise ParseError(f"解析失败 [{ext}]: {e}") from e

    # 截断保护
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[...已截断，原始长度 {len(text)} 字符]"
    return text.strip()


def _parse_pdf(file_path: Path) -> str:
    """PDF解析：优先 pypdf，失败回退 PyMuPDF"""
    # 优先 PyMuPDF (fitz) — 通常质量更好
    try:
        import fitz  # type: ignore

        doc = fitz.open(file_path)
        parts: list[str] = []
        for page in doc:
            parts.append(page.get_text())
        doc.close()
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass  # 回退到 pypdf

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        parts = [p.extract_text() or "" for p in reader.pages]
        return "\n".join(parts).strip()
    except Exception as e:
        raise ParseError(f"PDF解析失败: {e}") from e


def _parse_docx(file_path: Path) -> str:
    """Word .docx 解析：提取段落 + 表格"""
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []

        # 段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

        return "\n".join(parts).strip()
    except Exception as e:
        raise ParseError(f"DOCX解析失败: {e}") from e


def parse_file_with_pages(file_path: Path, max_chars: int = 50000) -> list[dict]:
    """解析文件，返回按页分割的文本块列表

    返回: [{"page_number": 1, "text": "..."}, ...]
    - PDF: 真实按页分割
    - DOCX/TXT: 按约 3000 字符分块模拟页码
    """
    ext = file_path.suffix.lower()
    try:
        if ext == ".pdf":
            return _parse_pdf_pages(file_path, max_chars)
        elif ext == ".docx":
            text = _parse_docx(file_path)
            return _chunk_text(text, max_chars)
        elif ext in (".txt", ".md"):
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return _chunk_text(text, max_chars)
        elif ext == ".doc":
            raise ParseError(".doc 格式不支持，请转换为 .docx 或 .pdf 后上传")
        else:
            raise ParseError(f"不支持的文件格式: {ext}")
    except ParseError:
        raise
    except Exception as e:
        raise ParseError(f"按页解析失败 [{ext}]: {e}") from e


def _roman_to_int(s: str) -> int | None:
    """罗马数字转阿拉伯数字，失败返回 None"""
    roman_map = {
        'M': 1000, 'CM': 900, 'D': 500, 'CD': 400,
        'C': 100, 'XC': 90, 'L': 50, 'XL': 40,
        'X': 10, 'IX': 9, 'V': 5, 'IV': 4, 'I': 1,
    }
    # 去除空格，如 "X X V" → "XXV"
    s = s.replace(' ', '').upper().strip()
    if not s:
        return None
    # 仅允许罗马数字字符
    if not all(c in 'IVXLCDM' for c in s):
        return None
    i = 0
    result = 0
    while i < len(s):
        for ch in ['CM', 'CD', 'XC', 'XL', 'IX', 'IV', 'M', 'D', 'C', 'L', 'X', 'V', 'I']:
            if s[i:].startswith(ch):
                result += roman_map[ch]
                i += len(ch)
                break
        else:
            return None
    return result


def _detect_page_number_from_text(text: str) -> int | None:
    """从纯文本片段中检测数字页码（支持多种格式）"""
    import re
    text = text.strip()
    if not text:
        return None

    # 1. 纯阿拉伯数字（如 "2", "24", "56"）
    if text.isdigit():
        return int(text)

    # 2. 带 | 分隔符：如 "24  | 第1 章" 或 "基础：逻辑和证明  | 25"
    if '|' in text:
        parts = [p.strip() for p in text.split('|')]
        for part in parts:
            cleaned = re.sub(r'^[\s\-–—]+|[\s\-–—]+$', '', part)
            if cleaned.isdigit():
                return int(cleaned)

    # 3. 罗马数字（如 "IX", "XII"）
    roman_val = _roman_to_int(text)
    if roman_val is not None:
        return roman_val

    return None


def _get_first_text_line(text: str) -> str:
    """获取文本块的第一行（换行符之前的部分）"""
    idx = text.find('\n')
    return text[:idx].strip() if idx >= 0 else text.strip()


def _detect_page_number_from_blocks(page) -> int | None:
    """利用 PyMuPDF 文本块坐标，从页面检测实际印刷页码

    策略（按优先级）：
    1. 四角检测：从页面四个角落（靠近纸张边缘）检测孤立的数字文本块
    2. 顶部边缘检测：从页面顶部附近检测数字文本块（页码常见于顶部左右两侧）
    3. 首行回退：如果以上均失败，从首行文本提取数字
    """
    try:
        blocks = page.get_text('blocks')
    except Exception:
        return None

    if not blocks:
        return None

    rect = page.rect
    page_w = rect.width
    page_h = rect.height

    import re

    # === 策略1：四角检测 ===
    # 角落区域：页面宽/高的 10% × 8% 以内
    corner_margin_x = page_w * 0.10
    corner_margin_y = page_h * 0.08

    candidates: list[tuple[int, float]] = []  # (page_number, corner_dist)

    for b in blocks:
        x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
        text = text.strip()
        if not text:
            continue

        # 取文本块的第一行（换行符前的内容）
        first_line = _get_first_text_line(text)
        if len(first_line) > 6:
            continue

        cx = (x0 + x1) / 2
        cy = (y0 + y1) / 2

        is_left = cx < corner_margin_x
        is_right = cx > page_w - corner_margin_x
        is_top = cy < corner_margin_y
        is_bottom = cy > page_h - corner_margin_y

        if not ((is_left or is_right) and (is_top or is_bottom)):
            continue

        pn = _detect_page_number_from_text(first_line)
        if pn is not None:
            corner_dist = min(
                ((cx - 0) ** 2 + (cy - 0) ** 2) ** 0.5,
                ((cx - page_w) ** 2 + (cy - 0) ** 2) ** 0.5,
                ((cx - 0) ** 2 + (cy - page_h) ** 2) ** 0.5,
                ((cx - page_w) ** 2 + (cy - page_h) ** 2) ** 0.5,
            )
            candidates.append((pn, corner_dist))

    if candidates:
        candidates.sort(key=lambda x: x[1])
        return candidates[0][0]

    # === 策略2：顶部边缘检测 ===
    # 教材页码常出现在页面顶部两侧（如 "22\n第1 章" 或 "基础：逻辑和证明\n23"）
    # 检测顶部 15% 区域内的文本块，提取第一行或最后一行的数字
    top_margin_y = page_h * 0.15
    side_margin_x = page_w * 0.15

    for b in blocks:
        x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
        text = text.strip()
        if not text:
            continue

        cy = (y0 + y1) / 2
        if cy > top_margin_y:
            continue

        cx = (x0 + x1) / 2
        # 在顶部但不在中央（页码通常在左右两侧）
        if side_margin_x < cx < page_w - side_margin_x:
            continue

        # 取第一行检测
        first_line = _get_first_text_line(text)
        if len(first_line) <= 6:
            pn = _detect_page_number_from_text(first_line)
            if pn is not None:
                return pn

        # 也检测最后一行（如 "基础：逻辑和证明\n23" 中 "23" 在最后一行）
        last_newline = text.rfind('\n')
        if last_newline >= 0:
            last_line = text[last_newline + 1:].strip()
            if last_line and len(last_line) <= 6:
                pn = _detect_page_number_from_text(last_line)
                if pn is not None:
                    return pn

    # === 策略3：从首行文本提取数字 ===
    # 某些页面格式如 "24"（单独一行就是页码）
    full_text = page.get_text().strip()
    lines = [l.strip() for l in full_text.split('\n') if l.strip()]
    if lines:
        first = lines[0]
        # 首行就是纯数字
        if first.isdigit() and len(first) <= 3:
            return int(first)
        # 尝试匹配开头的数字：如 "22 第1 章" → 22
        m = re.match(r'^(\d{1,3})\s', first)
        if m:
            return int(m.group(1))

    return None


def _infer_page_numbers(pages: list[dict]) -> list[dict]:
    """推断缺失的页码，基于已知页码进行线性插值"""
    import re

    # 第一步：为每页记录检测到的页码
    n = len(pages)

    # 第二步：前向填充——从已知页码推断后续页
    last_known = None
    last_known_idx = -1
    for i in range(n):
        pn = pages[i].get("_detected_page")
        if pn is not None:
            last_known = pn
            last_known_idx = i
            pages[i]["page_number"] = pn
        elif last_known is not None:
            # 从上一个已知页码推断
            inferred = last_known + (i - last_known_idx)
            pages[i]["page_number"] = inferred

    # 第三步：后向填充——如果前面还有未知的，从后往前填充
    next_known = None
    next_known_idx = -1
    for i in range(n - 1, -1, -1):
        pn = pages[i].get("_detected_page")
        if pn is not None:
            next_known = pn
            next_known_idx = i
        elif next_known is not None and pages[i].get("page_number") is None:
            inferred = next_known - (next_known_idx - i)
            if inferred > 0:
                pages[i]["page_number"] = inferred

    # 第四步：仍然缺失的，使用 PDF 索引 + 1
    for i in range(n):
        if pages[i].get("page_number") is None:
            pages[i]["page_number"] = i + 1

    return pages


def _parse_pdf_pages(file_path: Path, max_chars: int) -> list[dict]:
    """PDF 按页解析，返回每页的文本（检测实际印刷页码）"""
    try:
        import fitz  # type: ignore

        doc = fitz.open(file_path)
        raw_pages: list[dict] = []
        total_chars = 0
        for i in range(len(doc)):
            page = doc[i]
            text = page.get_text().strip()
            if not text:
                continue
            if total_chars + len(text) > max_chars:
                break
            # 从页面四个角落检测实际印刷页码
            detected = _detect_page_number_from_blocks(page)
            raw_pages.append({
                "_detected_page": detected,
                "text": text,
                "_pdf_idx": i,
            })
            total_chars += len(text)
        doc.close()

        if raw_pages:
            return _infer_page_numbers(raw_pages)

    except Exception:
        pass

    # 回退：pypdf 无坐标信息，从首行文本检测
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        raw_pages = []
        total_chars = 0
        for i, p in enumerate(reader.pages):
            text = (p.extract_text() or "").strip()
            if not text:
                continue
            if total_chars + len(text) > max_chars:
                break
            lines = [l.strip() for l in text.split('\n') if l.strip()]
            detected = _detect_page_number_from_text(lines[0]) if lines else None
            raw_pages.append({
                "_detected_page": detected,
                "text": text,
                "_pdf_idx": i,
            })
            total_chars += len(text)

        if raw_pages:
            return _infer_page_numbers(raw_pages)
        return raw_pages
    except Exception as e:
        raise ParseError(f"PDF按页解析失败: {e}") from e


def _chunk_text(text: str, max_chars: int, chunk_size: int = 3000) -> list[dict]:
    """将连续文本按约 3000 字符分块，模拟页码"""
    text = text.strip()
    if not text:
        return []
    pages = []
    total = 0
    page_num = 1
    while total < len(text) and total < max_chars:
        chunk = text[total: total + chunk_size]
        pages.append({"page_number": page_num, "text": chunk.strip()})
        total += chunk_size
        page_num += 1
    return pages


def make_preview(text: str, preview_chars: int = 500) -> str:
    """生成文本预览"""
    text = text.strip()
    if len(text) <= preview_chars:
        return text
    return text[:preview_chars] + "..."


def detect_material_type(filename: str, content: str) -> str:
    """启发式判断材料类型（六类枚举）
    syllabus:课程标准/大纲、textbook:教科书、reference:教参教辅、
    exercise_book:练习题册、paper:学术论文、other:其他
    """
    name = filename.lower()
    text_head = (content or "")[:500]
    # 课程标准/大纲
    if any(k in filename for k in ["大纲", "课程标准", "课标"]) or "syllabus" in name:
        return "syllabus"
    if any(k in filename for k in ["人才培养", "培养方案", "教学计划"]):
        return "syllabus"
    # 练习题册
    if any(k in filename for k in ["练习", "习题", "试题", "试卷", "题库"]):
        return "exercise_book"
    if any(k in text_head for k in ["一、选择题", "二、填空题", "单选题", "多选题", "简答题"]):
        return "exercise_book"
    # 学术论文
    if any(k in filename for k in ["论文", "期刊", "研究"]):
        return "paper"
    if "摘要" in text_head and "关键词" in text_head:
        return "paper"
    # 教参教辅（讲义/课件/教参/教案集）
    if any(k in filename for k in ["讲义", "课件", "教参", "教辅", "参考", "handout"]):
        return "reference"
    # 教科书
    if any(k in filename for k in ["教材", "教科书", "课本", "textbook"]):
        return "textbook"
    # 内容启发：有"第X章"+"节"结构偏向教材
    import re
    if re.search(r"第[一二三四五六七八九十百零0-9]+[章节]", content or "") and len(content or "") > 5000:
        return "textbook"
    return "other"
