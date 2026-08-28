"""教案模板的 .docx 解析与生成

提供两个核心函数：
- parse_template_docx(file_bytes) : 从上传的 .docx 解析出 structure_json
- generate_template_docx(structure_json, name) : 从 structure_json 生成可编辑的 .docx 字节流
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm




def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    tc_pr.append(shd)


def _add_table(doc, rows, cols, style="Table Grid"):
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = style
    return table


def _set_cell_text(cell, text, font_size=11, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.alignment = alignment
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _section_title(doc, text: str, color=(0x1c, 0x52, 0x47)):
    """添加居中章节标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    return p


def _make_table_row(table, label, value, label_color=(0x23, 0x66, 0x58), label_bg="D9EBE6"):
    """添加一行两列表格"""
    row = table.add_row()
    _set_cell_text(row.cells[0], label, font_size=10, bold=True, color=label_color)
    _set_cell_shading(row.cells[0], label_bg)
    _set_cell_text(row.cells[1], value, font_size=10)
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(12)
    return row


# ============================================================
#  解析：从 .docx 提取 structure_json
# ============================================================

def parse_template_docx(file_bytes: bytes) -> dict:
    """解析上传的 .docx 教案模板文件，提取 structure_json

    从 Word 文档中提取：
    - 课程名称、章节、总课时
    - 教学目标（知识/能力/素质）
    - 教学重难点
    - 教学过程阶段
    - 板书设计、作业、反思
    - 知识点出处
    """
    doc = Document(io.BytesIO(file_bytes))

    defaults: dict[str, Any] = {
        "course_name": "",
        "chapter": "",
        "total_minutes": 90,
        "knowledge_goal": "",
        "ability_goal": "",
        "value_goal": "",
        "key_points": [],
        "difficult_points": [],
        "difficult_strategy": "",
        "stages": [],
        "board_design": "",
        "homework": [],
        "reflection": "（课后填写）",
        "knowledge_sources": [],
    }

    # 收集所有文本
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    # 收集所有表格
    tables = doc.tables

    # 1. 从第一个表格提取标题信息
    title_text = ""
    subtitle_text = ""
    for table in tables:
        if table.rows and table.rows[0].cells:
            cell_text = table.rows[0].cells[0].text.strip()
            if cell_text and ("教案" in cell_text or "课程" in cell_text or "章" in cell_text):
                title_text = cell_text
                if len(table.rows) > 1:
                    subtitle_text = table.rows[1].cells[0].text.strip()
                break

    # 尝试从标题中解析课程名和章节
    title_match = re.search(r"(.+?)\s*[··]\s*(.+?)\s*(教案|教学设计)", title_text)
    if title_match:
        defaults["course_name"] = title_match.group(1).strip()
        defaults["chapter"] = title_match.group(2).strip()
    else:
        defaults["course_name"] = title_text.split("教案")[0].strip()[:50] if "教案" in title_text else title_text[:50]
        defaults["chapter"] = ""

    # 从副标题提取总课时
    if subtitle_text:
        time_match = re.search(r"(\d+)\s*分钟", subtitle_text)
        if time_match:
            defaults["total_minutes"] = int(time_match.group(1))

    # 2. 遍历表格提取结构化数据
    current_section = ""
    for table in tables:
        if not table.rows:
            continue
        first_cell_text = table.rows[0].cells[0].text.strip() if table.rows[0].cells else ""

        # 检测章节标题行（合并单元格，包含"一、""二、"等）
        if first_cell_text and re.match(r"[一二三四五六七八九十]、", first_cell_text):
            current_section = first_cell_text
            continue

        # 教学基本信息表
        if any(kw in first_cell_text for kw in ["课程名称", "授课章节", "授课对象", "课时"]):
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    label = cells[0].text.strip()
                    value = cells[1].text.strip()
                    if "课程名称" in label:
                        defaults["course_name"] = value or defaults["course_name"]
                    elif "授课章节" in label or "章节" in label:
                        defaults["chapter"] = value or defaults["chapter"]
                    elif "课时" in label or "分钟" in label:
                        tm = re.search(r"(\d+)", value)
                        if tm:
                            defaults["total_minutes"] = int(tm.group(1))

        # 教学目标表
        elif any(kw in first_cell_text for kw in ["知识目标", "能力目标", "素质目标", "思政目标", "教学目标"]):
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    label = cells[0].text.strip()
                    value = cells[1].text.strip()
                    if "知识目标" in label:
                        defaults["knowledge_goal"] = value
                    elif "能力目标" in label:
                        defaults["ability_goal"] = value
                    elif "素质目标" in label or "思政目标" in label or "价值目标" in label:
                        defaults["value_goal"] = value

        # 教学重难表
        elif any(kw in first_cell_text for kw in ["教学重点", "教学难点", "突破策略", "重难点"]):
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    label = cells[0].text.strip()
                    value = cells[1].text.strip()
                    if "重点" in label:
                        points = [p.strip() for p in re.split(r"[；;。\n]", value) if p.strip()]
                        defaults["key_points"] = points
                    elif "难点" in label:
                        points = [p.strip() for p in re.split(r"[；;。\n]", value) if p.strip()]
                        defaults["difficult_points"] = points
                    elif "突破" in label or "策略" in label:
                        defaults["difficult_strategy"] = value

        # 教学过程表（阶段/时长/教师行为/学生行为/设计意图/教学内容）
        elif any(kw in first_cell_text for kw in ["阶段", "时长", "教师行为", "学生行为", "教学过程"]):
            # 表头行，数据行从第二行开始
            for row_idx in range(1, len(table.rows)):
                cells = row.cells
                if len(cells) < 4:
                    continue
                stage_name = cells[0].text.strip()
                if not stage_name or "阶段" in stage_name:
                    continue
                # 解析阶段名和时长：如 "1. 导入\n5分钟"
                duration_min = 10
                clean_name = stage_name
                time_match = re.search(r"(\d+)\s*分钟", stage_name)
                if time_match:
                    duration_min = int(time_match.group(1))
                    clean_name = re.sub(r"[\d\.\s]*分钟", "", stage_name).strip()
                # 去掉前导编号
                clean_name = re.sub(r"^\d+[\.\s、]", "", clean_name).strip()

                stage = {
                    "name": clean_name or f"阶段{row_idx}",
                    "duration_min": duration_min,
                    "teacher_activity": cells[1].text.strip() if len(cells) > 1 else "",
                    "student_activity": cells[2].text.strip() if len(cells) > 2 else "",
                    "design_intent": cells[3].text.strip() if len(cells) > 3 else "",
                    "content": cells[4].text.strip() if len(cells) > 4 else "",
                }
                defaults["stages"].append(stage)

        # 板书设计
        elif "板书" in first_cell_text:
            if len(table.rows[0].cells) >= 2:
                defaults["board_design"] = table.rows[0].cells[1].text.strip()
            else:
                defaults["board_design"] = table.rows[0].cells[0].text.strip()

        # 课后作业
        elif "作业" in first_cell_text:
            homework_list = []
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    hw_text = cells[1].text.strip()
                    if hw_text:
                        homework_list.append(hw_text)
            defaults["homework"] = homework_list

        # 教学反思
        elif "反思" in first_cell_text:
            if len(table.rows[0].cells) >= 2:
                defaults["reflection"] = table.rows[0].cells[1].text.strip()
            else:
                defaults["reflection"] = table.rows[0].cells[0].text.strip()

    # 3. 从段落中提取未捕获的信息
    if not defaults["knowledge_goal"]:
        kg = re.search(r"(?:知识目标|教学目标)\s*[：:]\s*(.+?)(?:[。；;]|\n\n)", full_text)
        if kg:
            defaults["knowledge_goal"] = kg.group(1).strip()
    if not defaults["ability_goal"]:
        ag = re.search(r"(?:能力目标|技能目标)\s*[：:]\s*(.+?)(?:[。；;]|\n\n)", full_text)
        if ag:
            defaults["ability_goal"] = ag.group(1).strip()
    if not defaults["value_goal"]:
        vg = re.search(r"(?:素质目标|思政目标|价值目标|情感目标)\s*[：:]\s*(.+?)(?:[。；;]|\n\n)", full_text)
        if vg:
            defaults["value_goal"] = vg.group(1).strip()

    return {"defaults": defaults}


# ============================================================
#  生成：从 structure_json 生成可编辑的 .docx
# ============================================================

def generate_template_docx(structure_json: dict, name: str = "教案模板") -> bytes:
    """从模板结构生成可编辑的 .docx 文档

    生成的文档包含：
    - 标题：模板名称 + 说明
    - 表格化的教案结构预览（用户可看到模板包含哪些字段）
    - 一份空白骨架，用户可以直接在 Word 中编辑后上传
    """
    defaults = structure_json.get("defaults") or {}

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ========== 1. 标题 ==========
    title_table = _add_table(doc, rows=2, cols=2)
    title_cells = title_table.rows[0].cells
    title_cells[0].merge(title_cells[1])
    _set_cell_text(title_cells[0], f"教案模板 · {name}", font_size=18, bold=True,
                   color=(0x1c, 0x52, 0x47), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    sub_cells = title_table.rows[1].cells
    sub_cells[0].merge(sub_cells[1])
    _set_cell_text(sub_cells[0], "请在此模板中填写您的教案内容，保存后上传至备课助手即可使用",
                   font_size=10, bold=False, color=(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ========== 2. 使用说明 ==========
    _section_title(doc, "📋 使用说明")
    instructions = [
        "1. 请在下方的表格中填写您的教案内容（可直接在 Word 中编辑）",
        "2. 各字段说明：加粗列为字段名，右侧为填写内容",
        "3. 教学过程阶段可根据需要增删行（右键 → 插入行）",
        "4. 编辑完成后，保存文件并上传至备课助手即可作为教案模板使用",
        "5. 上传后系统会自动识别表格中的教学内容",
    ]
    for line in instructions:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ========== 3. 教学基本信息 ==========
    _section_title(doc, "一、教学基本信息")
    info_table = _add_table(doc, rows=5, cols=2)
    info_data = [
        ("课程名称", defaults.get("course_name", "") or ""),
        ("授课章节", defaults.get("chapter", "") or ""),
        ("授课对象", defaults.get("teaching_object", "") or ""),
        ("授课教师", defaults.get("teacher_name", "") or ""),
        ("课时安排", f'{defaults.get("total_minutes", 90)} 分钟'),
    ]
    for i, (label, value) in enumerate(info_data):
        _set_cell_text(info_table.cell(i, 0), label, font_size=10, bold=True, color=(0x23, 0x66, 0x58))
        _set_cell_shading(info_table.cell(i, 0), "D9EBE6")
        _set_cell_text(info_table.cell(i, 1), value, font_size=10)
    for row in info_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    # ========== 4. 教学目标 ==========
    _section_title(doc, "二、教学目标")
    goal_table = _add_table(doc, rows=3, cols=2)
    goal_data = [
        ("知识目标", defaults.get("knowledge_goal", "") or "", (0x63, 0x66, 0xb0), "E5E7F2"),
        ("能力目标", defaults.get("ability_goal", "") or "", (0xc7, 0x5c, 0x2e), "FDEBE6"),
        ("素质/思政目标", defaults.get("value_goal", "") or "", (0x2e, 0x7d, 0x6e), "D9EBE6"),
    ]
    for i, (label, value, color, bg) in enumerate(goal_data):
        _set_cell_text(goal_table.cell(i, 0), label, font_size=10, bold=True, color=color, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(goal_table.cell(i, 0), bg)
        _set_cell_text(goal_table.cell(i, 1), value, font_size=10)
    for row in goal_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    # ========== 5. 教学重难点 ==========
    _section_title(doc, "三、教学重难点")
    key_points = defaults.get("key_points") or []
    difficult_points = defaults.get("difficult_points") or []
    key_rows_count = 0
    if key_points:
        key_rows_count += 1
    if difficult_points:
        key_rows_count += 1
    if defaults.get("difficult_strategy"):
        key_rows_count += 1
    if key_rows_count == 0:
        key_rows_count = 1

    key_table = _add_table(doc, rows=key_rows_count, cols=2)
    row_idx = 0
    if key_points:
        _set_cell_text(key_table.cell(row_idx, 0), "教学重点", font_size=10, bold=True, color=(0xc7, 0x5c, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(key_table.cell(row_idx, 0), "FDEBE6")
        _set_cell_text(key_table.cell(row_idx, 1), "\n".join(f"{i}. {p}" for i, p in enumerate(key_points, 1)), font_size=10)
        row_idx += 1
    if difficult_points:
        _set_cell_text(key_table.cell(row_idx, 0), "教学难点", font_size=10, bold=True, color=(0xc7, 0x5c, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(key_table.cell(row_idx, 0), "FDEBE6")
        _set_cell_text(key_table.cell(row_idx, 1), "\n".join(f"{i}. {p}" for i, p in enumerate(difficult_points, 1)), font_size=10)
        row_idx += 1
    if defaults.get("difficult_strategy"):
        _set_cell_text(key_table.cell(row_idx, 0), "突破策略", font_size=10, bold=True, color=(0x2e, 0x7d, 0x6e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(key_table.cell(row_idx, 0), "D9EBE6")
        _set_cell_text(key_table.cell(row_idx, 1), defaults["difficult_strategy"], font_size=10)
    for row in key_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    # ========== 6. 教学过程设计 ==========
    _section_title(doc, "四、教学过程设计")
    stages = defaults.get("stages") or []

    # 至少显示一行空白行作为模板骨架
    display_stages = stages if stages else [{"name": "（请填写阶段名称）", "duration_min": 10, "teacher_activity": "", "student_activity": "", "design_intent": "", "content": ""}]

    headers = ["阶段/时长", "教师行为", "学生行为", "设计意图", "教学内容"]
    stages_table = _add_table(doc, rows=1, cols=len(headers))
    hdr_cells = stages_table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr_cells[i], h, font_size=10, bold=True, color=(0xff, 0xff, 0xff), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(hdr_cells[i], "2e7d6e")

    for i, stage in enumerate(display_stages, 1):
        row_cells = stages_table.add_row().cells
        stage_name = stage.get("name", f"阶段{i}")
        stage_duration = stage.get("duration_min", 10)
        _set_cell_text(row_cells[0], f"{i}. {stage_name}\n{stage_duration}分钟", font_size=10, bold=True, color=(0x23, 0x66, 0x58))
        _set_cell_text(row_cells[1], stage.get("teacher_activity", "") or "-", font_size=10)
        _set_cell_text(row_cells[2], stage.get("student_activity", "") or "-", font_size=10)
        _set_cell_text(row_cells[3], stage.get("design_intent", "") or "-", font_size=10)
        _set_cell_text(row_cells[4], stage.get("content", "") or "-", font_size=10)
        if i % 2 == 0:
            for cell in row_cells:
                _set_cell_shading(cell, "F0F7F5")

    doc.add_paragraph()

    # ========== 7. 板书设计 ==========
    _section_title(doc, "五、板书设计")
    board_table = _add_table(doc, rows=1, cols=2)
    _set_cell_text(board_table.cell(0, 0), "板书设计", font_size=10, bold=True, color=(0x23, 0x66, 0x58))
    _set_cell_shading(board_table.cell(0, 0), "D9EBE6")
    _set_cell_text(board_table.cell(0, 1), defaults.get("board_design", "") or "", font_size=10)
    board_table.cell(0, 0).width = Cm(3.5)
    board_table.cell(0, 1).width = Cm(12)
    doc.add_paragraph()

    # ========== 8. 课后作业 ==========
    _section_title(doc, "六、课后作业")
    homework = defaults.get("homework") or []
    display_hw = homework if homework else ["（请填写作业内容）"]
    hw_table = _add_table(doc, rows=len(display_hw), cols=2)
    for i, hw in enumerate(display_hw, 1):
        _set_cell_text(hw_table.cell(i-1, 0), str(i), font_size=10, bold=True, color=(0x23, 0x66, 0x58), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(hw_table.cell(i-1, 0), "F0F7F5")
        _set_cell_text(hw_table.cell(i-1, 1), hw, font_size=10)
    for row in hw_table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(14.5)
    doc.add_paragraph()

    # ========== 9. 教学反思 ==========
    _section_title(doc, "七、教学反思")
    ref_table = _add_table(doc, rows=1, cols=2)
    _set_cell_text(ref_table.cell(0, 0), "教学反思", font_size=10, bold=True, color=(0x23, 0x66, 0x58))
    _set_cell_shading(ref_table.cell(0, 0), "D9EBE6")
    _set_cell_text(ref_table.cell(0, 1), defaults.get("reflection", "") or "（课后填写）", font_size=10)
    ref_table.cell(0, 0).width = Cm(3.5)
    ref_table.cell(0, 1).width = Cm(12)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("本模板由备课助手智能体生成，可在 Word 中直接编辑后上传。")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()