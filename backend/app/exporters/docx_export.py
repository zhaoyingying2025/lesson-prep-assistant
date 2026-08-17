"""DOCX 导出 - 全表格形式"""
from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from ..models.schemas import LessonPlan


def _set_cell_shading(cell, color_hex: str) -> None:
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    tc_pr.append(shd)


def _add_table(doc, rows, cols, style="Table Grid"):
    """创建表格并设置边框"""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = style
    return table


def _set_cell_text(cell, text, font_size=11, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本样式"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.alignment = alignment
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def export_docx(plan: LessonPlan) -> bytes:
    """将教案导出为 DOCX 字节流（全表格形式）"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ========== 1. 标题表格 ==========
    title_table = _add_table(doc, rows=2, cols=2)
    # 标题行（合并两列）
    title_cells = title_table.rows[0].cells
    title_cells[0].merge(title_cells[1])
    _set_cell_text(title_cells[0], f"{plan.course_name} · {plan.chapter} 教案", font_size=18, bold=True,
                   color=(0x1c, 0x52, 0x47), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # 副标题行（合并两列）
    sub_cells = title_table.rows[1].cells
    sub_cells[0].merge(sub_cells[1])
    _set_cell_text(sub_cells[0], f"总课时：{plan.total_minutes} 分钟", font_size=10, bold=False,
                   color=(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # 标题表格边框加粗
    for row in title_table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn('w:tcBorders'))
            if tc_borders is None:
                tc_borders = tc_pr.makeelement(qn('w:tcBorders'))
                tc_pr.append(tc_borders)
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = tc_borders.makeelement(qn(f'w:{border_name}'))
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')
                border.set(qn('w:color'), '2e7d6e')
                tc_borders.append(border)

    doc.add_paragraph()

    # ========== 2. 教学基本信息（两列表格）==========
    info_title = doc.add_paragraph()
    info_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_title.add_run("一、教学基本信息")
    info_run.font.size = Pt(14)
    info_run.font.bold = True
    info_run.font.color.rgb = RGBColor(0x1c, 0x52, 0x47)

    info_rows = [
        ("课程名称", plan.course_name or ""),
        ("授课章节", plan.chapter or ""),
    ]
    if plan.teaching_object:
        info_rows.append(("授课对象", plan.teaching_object))
    if plan.teacher_name:
        info_rows.append(("授课教师", plan.teacher_name))
    info_rows.append(("课时安排", f"{plan.total_minutes} 分钟"))

    info_table = _add_table(doc, rows=len(info_rows), cols=2)
    for i, (label, value) in enumerate(info_rows):
        _set_cell_text(info_table.cell(i, 0), label, font_size=10, bold=True,
                       color=(0x23, 0x66, 0x58))
        _set_cell_shading(info_table.cell(i, 0), "D9EBE6")
        _set_cell_text(info_table.cell(i, 1), value, font_size=10)

    # 设置列宽
    for row in info_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    # ========== 3. 教学目标（三行两列表格，带彩色标签）==========
    goal_title = doc.add_paragraph()
    goal_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    goal_run = goal_title.add_run("二、教学目标")
    goal_run.font.size = Pt(14)
    goal_run.font.bold = True
    goal_run.font.color.rgb = RGBColor(0x1c, 0x52, 0x47)

    goal_table = _add_table(doc, rows=3, cols=2)

    _set_cell_text(goal_table.cell(0, 0), "知识目标", font_size=10, bold=True,
                   color=(0x63, 0x66, 0xb0), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shading(goal_table.cell(0, 0), "E5E7F2")
    _set_cell_text(goal_table.cell(0, 1), plan.knowledge_goal or "", font_size=10)

    _set_cell_text(goal_table.cell(1, 0), "能力目标", font_size=10, bold=True,
                   color=(0xc7, 0x5c, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shading(goal_table.cell(1, 0), "FDEBE6")
    _set_cell_text(goal_table.cell(1, 1), plan.ability_goal or "", font_size=10)

    _set_cell_text(goal_table.cell(2, 0), "素质/思政目标", font_size=10, bold=True,
                   color=(0x2e, 0x7d, 0x6e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shading(goal_table.cell(2, 0), "D9EBE6")
    _set_cell_text(goal_table.cell(2, 1), plan.value_goal or "", font_size=10)

    for row in goal_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    # ========== 4. 教学重难点（两列表格）==========
    key_title = doc.add_paragraph()
    key_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    key_run = key_title.add_run("三、教学重难点")
    key_run.font.size = Pt(14)
    key_run.font.bold = True
    key_run.font.color.rgb = RGBColor(0x1c, 0x52, 0x47)

    key_rows = []
    if plan.key_points:
        key_rows.append(("教学重点", "\n".join(f"{i}. {kp}" for i, kp in enumerate(plan.key_points, 1))))
    if plan.difficult_points:
        key_rows.append(("教学难点", "\n".join(f"{i}. {dp}" for i, dp in enumerate(plan.difficult_points, 1))))
    if plan.difficult_strategy:
        key_rows.append(("突破策略", plan.difficult_strategy))

    if key_rows:
        key_table = _add_table(doc, rows=len(key_rows), cols=2)
        for i, (label, value) in enumerate(key_rows):
            _set_cell_text(key_table.cell(i, 0), label, font_size=10, bold=True,
                           color=(0xc7, 0x5c, 0x2e) if i < 2 else (0x2e, 0x7d, 0x6e),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(key_table.cell(i, 0), "FDEBE6" if i < 2 else "D9EBE6")
            _set_cell_text(key_table.cell(i, 1), value, font_size=10)
        for row in key_table.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(12.5)
    else:
        doc.add_paragraph("暂无", style="Body Text")

    doc.add_paragraph()

    # ========== 5. 教学过程设计（多列表格）==========
    stage_title = doc.add_paragraph()
    stage_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stage_run = stage_title.add_run("四、教学过程设计")
    stage_run.font.size = Pt(14)
    stage_run.font.bold = True
    stage_run.font.color.rgb = RGBColor(0x1c, 0x52, 0x47)

    if plan.stages:
        headers = ["阶段/时长", "教师行为", "学生行为", "设计意图", "教学内容"]
        stages_table = _add_table(doc, rows=1, cols=len(headers))

        # 表头
        hdr_cells = stages_table.rows[0].cells
        for i, h in enumerate(headers):
            _set_cell_text(hdr_cells[i], h, font_size=10, bold=True,
                           color=(0xff, 0xff, 0xff), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(hdr_cells[i], "2e7d6e")

        # 数据行
        for i, stage in enumerate(plan.stages, 1):
            row_cells = stages_table.add_row().cells
            _set_cell_text(row_cells[0], f"{i}. {stage.name}\n{stage.duration_min}分钟",
                           font_size=10, bold=True, color=(0x23, 0x66, 0x58))
            _set_cell_text(row_cells[1], stage.teacher_activity or "-", font_size=10)
            _set_cell_text(row_cells[2], stage.student_activity or "-", font_size=10)
            _set_cell_text(row_cells[3], stage.design_intent or "-", font_size=10)
            _set_cell_text(row_cells[4], stage.content or "-", font_size=10)

            # 隔行底色
            if i % 2 == 0:
                for cell in row_cells:
                    _set_cell_shading(cell, "F0F7F5")

        # 设置列宽
        col_widths = [Cm(2.2), Cm(3.2), Cm(3.2), Cm(3.2), Cm(3.2)]
        for row in stages_table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(col_widths):
                    cell.width = col_widths[idx]
    else:
        doc.add_paragraph("暂无教学过程", style="Body Text")

    doc.add_paragraph()

    # ========== 6. 板书设计（两列表格）==========
    if plan.board_design:
        board_title = doc.add_paragraph()
        board_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        board_run = board_title.add_run("五、板书设计")
        board_run.font.size = Pt(14)
        board_run.font.bold = True
        board_run.font.color.rgb = RGBColor(0x1c, 0x52, 0x47)

        board_table = _add_table(doc, rows=1, cols=2)
        _set_cell_text(board_table.cell(0, 0), "板书设计", font_size=10, bold=True,
                       color=(0x23, 0x66, 0x58))
        _set_cell_shading(board_table.cell(0, 0), "D9EBE6")
        _set_cell_text(board_table.cell(0, 1), plan.board_design, font_size=10)
        board_table.cell(0, 0).width = Cm(3)
        board_table.cell(0, 1).width = Cm(12)
        doc.add_paragraph()

    # ========== 7. 课后作业（两列表格）==========
    if plan.homework:
        hw_title = doc.add_paragraph()
        hw_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hw_run = hw_title.add_run("六、课后作业")
        hw_run.font.size = Pt(14)
        hw_run.font.bold = True
        hw_run.font.color.rgb = RGBColor(0x1c, 0x52, 0x47)

        hw_table = _add_table(doc, rows=len(plan.homework), cols=2)
        for i, hw in enumerate(plan.homework, 1):
            _set_cell_text(hw_table.cell(i-1, 0), str(i), font_size=10, bold=True,
                           color=(0x23, 0x66, 0x58), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(hw_table.cell(i-1, 0), "F0F7F5")
            _set_cell_text(hw_table.cell(i-1, 1), hw, font_size=10)
        for row in hw_table.rows:
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(14)
        doc.add_paragraph()

    # ========== 8. 教学反思（两列表格）==========
    ref_title = doc.add_paragraph()
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_title.add_run("七、教学反思")
    ref_run.font.size = Pt(14)
    ref_run.font.bold = True
    ref_run.font.color.rgb = RGBColor(0x1c, 0x52, 0x47)

    ref_table = _add_table(doc, rows=1, cols=2)
    _set_cell_text(ref_table.cell(0, 0), "教学反思", font_size=10, bold=True,
                   color=(0x23, 0x66, 0x58))
    _set_cell_shading(ref_table.cell(0, 0), "D9EBE6")
    _set_cell_text(ref_table.cell(0, 1), plan.reflection or "（课后填写）", font_size=10)
    ref_table.cell(0, 0).width = Cm(3)
    ref_table.cell(0, 1).width = Cm(12)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("本教案由备课助手智能体辅助生成，仅供教学参考。")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()